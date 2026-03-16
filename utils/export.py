"""Note export helpers: markdown with base64 images, PDF, and DOCX."""

from __future__ import annotations

import base64
import logging
import re
import zipfile
from collections import defaultdict
from io import BytesIO
from pathlib import Path
from typing import TYPE_CHECKING, Optional

if TYPE_CHECKING:
    from models.document import Block, Document

LOGGER = logging.getLogger(__name__)


# Matches section-number prefixes like "4.1.2" at the start of a text block.
_DOC_SEC_RE = re.compile(r"^\d+(?:\.\d+)+\s+\S")
# Extracts the first section number from any string (e.g. "## 4.1.2 Title" → "4.1.2").
_SEC_NUM_RE = re.compile(r"(\d+(?:\.\d+)+)")


def _build_section_page_ranges(
    doc: "Document",
    n_sections: int,
    section_headings: Optional[list[str]] = None,
) -> list[tuple[int, int]]:
    """Compute the page range covered by each note section.

    Primary strategy (when ``section_headings`` is provided):
        1. Extract section-header blocks from the document by matching the
           pattern ``<number> <title>`` (e.g. "4.1.2 Activation Functions").
        2. Build a map of section-number → start page.
        3. For each note section heading, extract its section number and look up
           the start page.  The end page is one before the next document section
           that starts on a later page.
        4. For unmatched headings, fall back to the even-split strategy below.

    Fallback strategy (even split):
        Divide all body blocks (non-image, with page) evenly across n_sections
        and return the min/max page of each chunk.

    Args:
        doc: Parsed Document.
        n_sections: Number of note sections.
        section_headings: List of raw note section headings (e.g. "## 4.1.2 …").
            When provided, enables the primary section-number matching strategy.

    Returns:
        List of (min_page, max_page) tuples, one per section.
    """
    # --- collect document section-header pages ---
    doc_sec_map: dict[str, int] = {}  # "4.1.2" → page
    if section_headings:
        for b in doc.blocks:
            if b.metadata.page and not b.image_path and _DOC_SEC_RE.match(b.content.strip()):
                m = _SEC_NUM_RE.match(b.content.strip())
                if m:
                    doc_sec_map[m.group(1)] = b.metadata.page

    # --- primary strategy: section-number matching ---
    if doc_sec_map and section_headings:
        max_page = max((b.metadata.page for b in doc.blocks if b.metadata.page), default=9999)
        # sorted (sec_num, page) for next-boundary lookup
        sec_pages_sorted = sorted(doc_sec_map.values())

        def _page_range_for_sec(sec_num: str) -> Optional[tuple[int, int]]:
            if sec_num not in doc_sec_map:
                return None
            lo = doc_sec_map[sec_num]
            # next document section that starts on a later page
            hi = next((p - 1 for p in sec_pages_sorted if p > lo), max_page)
            return (lo, max(lo, hi))

        # body blocks for fallback chunks
        body_blocks = sorted(
            [b for b in doc.blocks if b.metadata.page is not None and not b.image_path],
            key=lambda b: b.order,
        )
        total = len(body_blocks)

        ranges: list[tuple[int, int]] = []
        for i, heading in enumerate(section_headings):
            m = _SEC_NUM_RE.search(heading)
            if m:
                pr = _page_range_for_sec(m.group(1))
                if pr:
                    ranges.append(pr)
                    continue
            # fallback: even-split chunk for this index
            if total:
                start = (i * total) // n_sections
                end = ((i + 1) * total) // n_sections
                chunk = body_blocks[start:end]
                pages = [b.metadata.page for b in chunk if b.metadata.page is not None]
                if pages:
                    ranges.append((min(pages), max(pages)))
                    continue
            prev = ranges[-1][1] if ranges else 1
            ranges.append((prev, prev))
        return ranges

    # --- fallback: even split ---
    body_blocks = sorted(
        [b for b in doc.blocks if b.metadata.page is not None and not b.image_path],
        key=lambda b: b.order,
    )
    m_total = len(body_blocks)
    if m_total == 0:
        return [(1, 9999)] * n_sections
    ranges = []
    for i in range(n_sections):
        start = (i * m_total) // n_sections
        end = ((i + 1) * m_total) // n_sections
        chunk = body_blocks[start:end]
        pages = [b.metadata.page for b in chunk if b.metadata.page is not None]
        if pages:
            ranges.append((min(pages), max(pages)))
        else:
            prev = ranges[-1][1] if ranges else 1
            ranges.append((prev, prev))
    return ranges


def _place_figures_page_based(
    fig_blocks: list["Block"],
    section_page_ranges: list[tuple[int, int]],
) -> dict[int, list]:
    """Map each figure block to a section index based on its page number.

    1st pass: exact range match (lo <= page <= hi).
    2nd pass: nearest midpoint fallback.
    Figures without a page are appended to the last section.

    Returns:
        dict[section_idx → list[Block]]
    """
    section_figs: dict[int, list] = defaultdict(list)
    for b in fig_blocks:
        p = b.metadata.page
        if p is None:
            section_figs[len(section_page_ranges) - 1].append(b)
            continue
        # Use the *last* matching section (deepest / most specific) so that when a
        # parent heading and its first subsection share the same start page their
        # ranges overlap and we correctly assign the figure to the subsection.
        all_matches = [i for i, (lo, hi) in enumerate(section_page_ranges) if lo <= p <= hi]
        if all_matches:
            matched = all_matches[-1]
        else:
            matched = min(
                range(len(section_page_ranges)),
                key=lambda i: abs((section_page_ranges[i][0] + section_page_ranges[i][1]) / 2 - p),
            )
        section_figs[matched].append(b)
    return dict(section_figs)


def interleave_figures_into_sections(
    note_markdown: str,
    fig_blocks: list["Block"],
    doc: Optional["Document"] = None,
) -> list[dict]:
    """Split note into sections and interleave figures at UI-matching positions.

    When ``doc`` is provided, uses page-to-section mapping (preferred) which
    computes the actual page range of each section from document body blocks and
    places each figure in the section whose page range contains the figure's page.
    When ``doc`` is None, falls back to the legacy page interpolation logic for
    backward compatibility.

    Returns a list of items in render order::

        [
            {"type": "section", "markdown": "## 서론\\n텍스트..."},
            {"type": "figure", "block": <Block>, "caption": "그림 (page 3)"},
            {"type": "section", "markdown": "## 신경망\\n텍스트..."},
            ...
        ]
    """
    from llm.note_editor import _split_sections

    sections = _split_sections(note_markdown)
    n = len(sections)
    if n == 0:
        return [{"type": "section", "markdown": note_markdown}]

    # Build section markdown strings
    def _section_md(heading: str, body: str) -> str:
        if heading:
            return f"{heading}\n\n{body}".strip()
        return body

    if doc is not None:
        # Page-based mapping: use section-header page boundaries when available
        headings = [h for h, _ in sections]
        ranges = _build_section_page_ranges(doc, n, section_headings=headings)
        section_figs = _place_figures_page_based(fig_blocks, ranges)
    else:
        # Legacy page interpolation fallback
        pages = [b.metadata.page for b in fig_blocks]
        valid_pages = [p for p in pages if p is not None]
        page_min = min(valid_pages) if valid_pages else 1
        page_max = max(valid_pages) if valid_pages else 1
        page_range = max(page_max - page_min, 1)

        section_figs_dd: dict[int, list] = defaultdict(list)
        page_counters: dict = defaultdict(int)
        for b in fig_blocks:
            p = b.metadata.page if b.metadata.page is not None else page_max
            ratio = (p - page_min) / page_range
            base_idx = min(int(ratio * n), n - 1)
            count = page_counters[b.metadata.page]
            page_counters[b.metadata.page] += 1
            idx = min(base_idx + count, n - 1)
            section_figs_dd[idx].append(b)
        section_figs = dict(section_figs_dd)

    items: list[dict] = []
    for i, (heading, body) in enumerate(sections):
        # Mirror UI: skip heading-only sections with no body and no figures
        if heading and not body and not section_figs.get(i):
            continue
        md = _section_md(heading, body)
        if md:
            items.append({"type": "section", "markdown": md})
        for b in section_figs.get(i, []):
            if b.image_path and Path(b.image_path).exists():
                caption = b.metadata.caption or f"그림 (page {b.metadata.page})"
                items.append({"type": "figure", "block": b, "caption": caption})

    return items


def export_markdown_zip(
    note_markdown: str,
    title: str,
    fig_blocks: list["Block"],
    doc: Optional["Document"] = None,
) -> bytes:
    """Return note as a ZIP archive containing note.md + figures/ directory.

    Images are referenced via relative paths (``figures/{order}.png``) so the
    markdown renders correctly when the ZIP is extracted.  This avoids the
    bloated base64 inline approach and keeps the markdown file human-readable.

    Args:
        note_markdown: Raw note markdown (without title prefix).
        title: Document / note title.
        fig_blocks: Blocks with ``image_path`` set (pre-filtered by caller).
        doc: Parsed Document used for page-to-section mapping (preferred).
            Falls back to page interpolation when None.

    Returns:
        ZIP file bytes containing ``note.md`` and ``figures/*.png``.
    """
    items = interleave_figures_into_sections(note_markdown, fig_blocks, doc=doc)
    md_parts: list[str] = [f"# {title}\n\n"]
    image_files: list[tuple[str, bytes]] = []  # (path_in_zip, raw_bytes)

    for item in items:
        if item["type"] == "section":
            md_parts.append(item["markdown"] + "\n\n")
        elif item["type"] == "figure":
            b = item["block"]
            if not b.image_path or not Path(b.image_path).exists():
                continue
            fname = f"figures/{b.order}.png"
            caption = item["caption"]
            md_parts.append(f"![{caption}]({fname})\n\n")
            md_parts.append(f"*{caption}*\n\n")
            image_files.append((fname, Path(b.image_path).read_bytes()))

    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("note.md", "".join(md_parts))
        for path, data in image_files:
            zf.writestr(path, data)
    return buf.getvalue()


# Bundled Noto Sans KR font files (Regular + Bold) for Korean PDF rendering.
# Paths are relative to the project root; resolved at call time.
_FONT_DIR = Path(__file__).parent.parent / "data" / "fonts"
_FONT_REGULAR = _FONT_DIR / "NotoSansKR-Regular.ttf"
_FONT_BOLD = _FONT_DIR / "NotoSansKR-Bold.ttf"


def _korean_font_css() -> str:
    """Return @font-face CSS block if Noto Sans KR TTF files are present.

    Font files are expected at ``data/fonts/NotoSansKR-{Regular,Bold}.ttf``.
    Uses file:// URIs so xhtml2pdf can load fonts directly without base64
    encoding overhead.  Returns an empty string when files are not found so
    the function degrades gracefully (Korean text will render as boxes).
    """
    if not (_FONT_REGULAR.exists() and _FONT_BOLD.exists()):
        LOGGER.warning(
            "Noto Sans KR font files not found at %s — Korean PDF text may not render correctly",
            _FONT_DIR,
        )
        return ""

    reg_uri = _FONT_REGULAR.as_uri()
    bold_uri = _FONT_BOLD.as_uri()
    return (
        f"@font-face {{font-family:'NotoSansKR';font-weight:normal;"
        f"src:url('{reg_uri}') format('truetype');}}"
        f"@font-face {{font-family:'NotoSansKR';font-weight:bold;"
        f"src:url('{bold_uri}') format('truetype');}}"
    )


def export_pdf(
    note_markdown: str,
    title: str,
    fig_blocks: list["Block"],
    doc: Optional["Document"] = None,
) -> bytes:
    """Return note as PDF bytes with inline images at UI-matching positions.

    Requires ``xhtml2pdf`` (``pip install xhtml2pdf``).
    Korean text is rendered using Noto Sans KR if the TTF files are present at
    ``data/fonts/NotoSansKR-Regular.ttf`` and ``data/fonts/NotoSansKR-Bold.ttf``.

    Args:
        note_markdown: Raw note markdown (without title prefix).
        title: Document / note title shown as the PDF ``<h1>``.
        fig_blocks: Blocks with ``image_path`` set.
        doc: Parsed Document used for page-to-section mapping (preferred).
            Falls back to page interpolation when None.

    Returns:
        PDF file contents as bytes.

    Raises:
        RuntimeError: If xhtml2pdf reports errors during PDF generation.
    """
    import html as _html

    import markdown as md_lib
    from xhtml2pdf import pisa

    items = interleave_figures_into_sections(note_markdown, fig_blocks, doc=doc)

    html_parts: list[str] = []
    for item in items:
        if item["type"] == "section":
            html_parts.append(
                md_lib.markdown(item["markdown"], extensions=["fenced_code", "tables"])
            )
        elif item["type"] == "figure":
            b = item["block"]
            if not b.image_path:
                continue
            img_bytes = Path(b.image_path).read_bytes()
            b64 = base64.b64encode(img_bytes).decode()
            caption = item["caption"]
            html_parts.append(
                f'<div style="margin:20px 0;text-align:center;">'
                f'<img src="data:image/png;base64,{b64}" style="max-width:100%;" />'
                f'<p style="color:#666;font-size:0.9em;">{_html.escape(caption)}</p>'
                f"</div>"
            )

    font_css = _korean_font_css()
    body_font = "'NotoSansKR',Helvetica,Arial,sans-serif" if font_css else "Helvetica,Arial,sans-serif"

    full_html = (
        "<!DOCTYPE html>"
        '<html><head><meta charset="utf-8">'
        "<style>"
        f"{font_css}"
        f"body{{font-family:{body_font};max-width:800px;"
        "margin:0 auto;padding:20px;line-height:1.8;}"
        "h1{color:#C4553A;}"
        "h2{color:#333;border-bottom:1px solid #ddd;padding-bottom:8px;}"
        "code{background:#f4f4f4;padding:2px 6px;border-radius:3px;}"
        "pre{background:#f4f4f4;padding:12px;border-radius:6px;overflow-x:auto;}"
        "</style>"
        "</head><body>"
        f"<h1>{_html.escape(title)}</h1>"
        + "".join(html_parts)
        + "</body></html>"
    )

    buffer = BytesIO()
    result = pisa.CreatePDF(full_html, dest=buffer)
    if result.err:
        raise RuntimeError(f"xhtml2pdf error: {result.err}")
    return buffer.getvalue()


def export_docx(
    note_markdown: str,
    title: str,
    fig_blocks: list["Block"],
    doc: Optional["Document"] = None,
) -> bytes:
    """Return note as DOCX bytes with inline images at UI-matching positions.

    Requires ``python-docx`` (``pip install python-docx``).

    Args:
        note_markdown: Raw note markdown (without title prefix).
        title: Document / note title shown as the DOCX ``Heading 0``.
        fig_blocks: Blocks with ``image_path`` set.
        doc: Parsed Document used for page-to-section mapping (preferred).
            Falls back to page interpolation when None.

    Returns:
        DOCX file contents as bytes.
    """
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    docx_doc = DocxDocument()
    docx_doc.add_heading(title, level=0)

    items = interleave_figures_into_sections(note_markdown, fig_blocks, doc=doc)

    for item in items:
        if item["type"] == "section":
            for line in item["markdown"].split("\n"):
                stripped = line.strip()
                if not stripped:
                    continue
                if stripped.startswith("## "):
                    docx_doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("### "):
                    docx_doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith(("- ", "* ")):
                    docx_doc.add_paragraph(stripped[2:], style="List Bullet")
                else:
                    docx_doc.add_paragraph(stripped)

        elif item["type"] == "figure":
            b = item["block"]
            docx_doc.add_picture(b.image_path, width=Inches(5.5))
            caption_para = docx_doc.add_paragraph(item["caption"])
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if caption_para.runs:
                caption_para.runs[0].font.size = Pt(9)

    buffer = BytesIO()
    docx_doc.save(buffer)
    return buffer.getvalue()
